"""Document processing with tiktoken-based chunking."""
import asyncio
import logging
import os

import docx
import tiktoken
from langchain_community.document_loaders import PyPDFLoader, TextLoader
from langchain_community.vectorstores import Qdrant
from langchain_core.documents import Document

from app.core.config import settings
from app.core.metrics import DOCUMENT_PROCESSED
from app.infrastructure.qdrant import qdrant_service
from app.services.embeddings import embeddings_service

logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt"}


class TokenTextSplitter:
    """Split text by token count using tiktoken."""

    def __init__(
        self,
        chunk_size: int = 256,
        chunk_overlap: int = 100,
        encoding_name: str = "cl100k_base",
    ) -> None:
        """Initialize splitter with token limits."""
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.encoding = tiktoken.get_encoding(encoding_name)

    def _token_count(self, text: str) -> int:
        """Count tokens in text."""
        return len(self.encoding.encode(text))

    def _split_text(self, text: str) -> list[str]:
        """Split text into chunks by token count."""
        tokens = self.encoding.encode(text)
        chunks = []
        start = 0

        while start < len(tokens):
            end = min(start + self.chunk_size, len(tokens))
            chunk_tokens = tokens[start:end]
            chunk_text = self.encoding.decode(chunk_tokens)
            chunks.append(chunk_text)

            if end >= len(tokens):
                break

            start = end - self.chunk_overlap

        return chunks

    def split_documents(self, documents: list[Document]) -> list[Document]:
        """Split documents into token-sized chunks."""
        result = []
        for doc in documents:
            chunks = self._split_text(doc.page_content)
            for i, chunk in enumerate(chunks):
                metadata = doc.metadata.copy()
                metadata["chunk_index"] = i
                metadata["token_count"] = self._token_count(chunk)
                result.append(Document(page_content=chunk, metadata=metadata))
        return result


text_splitter = TokenTextSplitter(
    chunk_size=settings.CHUNK_SIZE_TOKENS,
    chunk_overlap=settings.CHUNK_OVERLAP_TOKENS,
    encoding_name=settings.TIKTOKEN_ENCODING,
)


def load_document(file_path: str) -> list[Document]:
    """Load document by file extension."""
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        return PyPDFLoader(file_path).load()

    if ext == ".docx":
        doc = docx.Document(file_path)
        text = "\n".join([p.text for p in doc.paragraphs])
        return [Document(page_content=text, metadata={"source": file_path})]

    if ext == ".txt":
        return TextLoader(file_path, encoding="utf-8").load()

    raise ValueError(f"Unsupported format: {ext}")


def split_documents(documents: list[Document]) -> list[Document]:
    """Split documents into token-sized chunks."""
    return text_splitter.split_documents(documents)


def upload_to_qdrant(chunks: list[Document]) -> None:
    """Upload document chunks to Qdrant."""
    logger.info(f"Uploading {len(chunks)} chunks to Qdrant")
    Qdrant.from_documents(
        chunks,
        embeddings_service.model,
        url=qdrant_service.url,
        prefer_grpc=False,
        collection_name=settings.QDRANT_COLLECTION,
    )


async def process_document(file_path: str) -> None:
    """Process document: load, split by tokens, embed, and store."""
    if not os.path.exists(file_path):
        DOCUMENT_PROCESSED.labels(status="not_found").inc()
        raise FileNotFoundError(f"File not found: {file_path}")

    ext = os.path.splitext(file_path)[1].lower()
    if ext not in SUPPORTED_EXTENSIONS:
        DOCUMENT_PROCESSED.labels(status="unsupported").inc()
        raise ValueError(f"Unsupported format: {ext}")

    loop = asyncio.get_running_loop()

    def process() -> None:
        documents = load_document(file_path)
        chunks = split_documents(documents)
        logger.info(f"Split into {len(chunks)} token-based chunks")
        upload_to_qdrant(chunks)

    try:
        await loop.run_in_executor(None, process)
        DOCUMENT_PROCESSED.labels(status="success").inc()
        logger.info(f"Document processed: {file_path}")
    except Exception as e:
        DOCUMENT_PROCESSED.labels(status="error").inc()
        logger.error(f"Document processing failed: {e}")
        raise
