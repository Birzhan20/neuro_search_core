"""RAG Service - Async gRPC server with RabbitMQ consumer."""
from __future__ import annotations

import asyncio
import json
import logging
import os
import sys
from typing import TYPE_CHECKING, Any

import grpc
from aio_pika import connect_robust
from aio_pika.abc import AbstractIncomingMessage
from dotenv import load_dotenv
import docx
from langchain_community.document_loaders import PyPDFLoader, TextLoader
from langchain_community.vectorstores import Qdrant
from langchain_core.messages import HumanMessage, SystemMessage
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_openai import ChatOpenAI
from langchain_text_splitters import RecursiveCharacterTextSplitter
from qdrant_client import AsyncQdrantClient
from qdrant_client.http import models

if TYPE_CHECKING:
    pass

load_dotenv()

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

current_dir = os.path.dirname(os.path.abspath(__file__))
project_root = os.path.dirname(current_dir)
proto_dir = os.path.join(project_root, "proto")
sys.path.insert(0, project_root)
sys.path.insert(0, proto_dir)

from proto import rag_service_pb2  # noqa: E402
from proto import rag_service_pb2_grpc  # noqa: E402

RABBITMQ_URL = os.getenv("RABBITMQ_URL", "amqp://guest:guest@rabbitmq:5672/")
QDRANT_HOST = os.getenv("QDRANT_HOST", "qdrant")
QDRANT_PORT = int(os.getenv("QDRANT_PORT", "6333"))
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
COLLECTION_NAME = "documents"

logger.info("Loading embeddings model...")
embeddings_model = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")

logger.info("Initializing LLM...")
llm = ChatOpenAI(temperature=0, model="gpt-4o-mini", api_key=OPENAI_API_KEY)


async def init_qdrant_collection() -> None:
    """Initialize Qdrant collection if it doesn't exist."""
    client = AsyncQdrantClient(host=QDRANT_HOST, port=QDRANT_PORT)
    try:
        collections = await client.get_collections()
        exists = any(c.name == COLLECTION_NAME for c in collections.collections)

        if not exists:
            logger.info("Creating collection %s", COLLECTION_NAME)
            await client.create_collection(
                collection_name=COLLECTION_NAME,
                vectors_config=models.VectorParams(
                    size=384,
                    distance=models.Distance.COSINE,
                ),
            )
    finally:
        await client.close()


def load_document(file_path: str) -> list[Any]:
    """Load document based on file extension."""
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        loader = PyPDFLoader(file_path)
        return loader.load()
    elif ext == ".docx":
        doc = docx.Document(file_path)
        text = "\n".join([para.text for para in doc.paragraphs])
        from langchain_core.documents import Document
        return [Document(page_content=text, metadata={"source": file_path})]
    elif ext == ".txt":
        loader = TextLoader(file_path, encoding="utf-8")
        return loader.load()
    else:
        raise ValueError(f"Unsupported file format: {ext}")


async def process_file_task(message: AbstractIncomingMessage) -> None:
    """Process incoming file task from RabbitMQ."""
    async with message.process():
        try:
            body = message.body.decode()
            data: dict[str, Any] = json.loads(body)
            task_id: str = data.get("task_id", "unknown")
            file_path: str = data.get("file_path", "")
            logger.info("Processing task %s: %s", task_id, file_path)

            if not os.path.exists(file_path):
                logger.error("File not found: %s", file_path)
                return

            ext = os.path.splitext(file_path)[1].lower()
            if ext not in (".pdf", ".docx", ".txt"):
                logger.error("Unsupported file format: %s", ext)
                return

            loop = asyncio.get_running_loop()

            def processing_logic() -> list[Any]:
                pages = load_document(file_path)
                text_splitter = RecursiveCharacterTextSplitter(
                    chunk_size=1000,
                    chunk_overlap=200,
                )
                return text_splitter.split_documents(pages)

            splits = await loop.run_in_executor(None, processing_logic)
            logger.info("Vectorizing %d chunks", len(splits))

            url = f"http://{QDRANT_HOST}:{QDRANT_PORT}"

            def upload_logic() -> None:
                Qdrant.from_documents(
                    splits,
                    embeddings_model,
                    url=url,
                    prefer_grpc=False,
                    collection_name=COLLECTION_NAME,
                )

            await loop.run_in_executor(None, upload_logic)
            logger.info("Task %s completed", task_id)

        except Exception as e:
            logger.exception("Error processing task: %s", e)


async def consume_rabbitmq() -> None:
    """Connect to RabbitMQ and consume messages."""
    retries = 5
    connection = None

    while retries > 0:
        try:
            connection = await connect_robust(RABBITMQ_URL)
            logger.info("Consumer connected to RabbitMQ")
            break
        except Exception:
            logger.warning("Waiting for RabbitMQ...")
            await asyncio.sleep(5)
            retries -= 1

    if not connection:
        logger.error("Could not connect to RabbitMQ")
        return

    async with connection:
        channel = await connection.channel()
        queue = await channel.declare_queue("ingestion_queue", durable=True)
        await channel.set_qos(prefetch_count=1)

        async with queue.iterator() as queue_iter:
            async for message in queue_iter:
                await process_file_task(message)


class RagService(rag_service_pb2_grpc.RagServiceServicer):
    """gRPC RAG Service implementation."""

    async def GetAnswer(
        self,
        request: rag_service_pb2.ChatRequest,
        context: grpc.aio.ServicerContext,
    ) -> rag_service_pb2.ChatResponse:
        """Handle chat request with RAG pipeline."""
        query = request.message
        logger.info("Query: %s", query)

        try:
            loop = asyncio.get_running_loop()
            query_vector: list[float] = await loop.run_in_executor(
                None,
                embeddings_model.embed_query,
                query,
            )

            client = AsyncQdrantClient(host=QDRANT_HOST, port=QDRANT_PORT)
            try:
                search_result = await client.query_points(
                    collection_name=COLLECTION_NAME,
                    query=query_vector,
                    limit=3,
                    with_payload=True,
                )
            finally:
                await client.close()

            context_text = ""
            sources_proto: list[rag_service_pb2.Source] = []

            for point in search_result.points:
                payload: dict[str, Any] = point.payload or {}
                doc_name = str(payload.get("source", "unknown"))
                page_num = int(payload.get("page", 0))
                text_chunk = str(payload.get("page_content", ""))
                score = float(point.score) if point.score else 0.0

                context_text += (
                    f"---\nDocument: {doc_name} (page {page_num})\n"
                    f"Text: {text_chunk}\n"
                )
                sources_proto.append(
                    rag_service_pb2.Source(
                        doc_name=os.path.basename(doc_name),
                        page=page_num,
                        score=score,
                    )
                )

            if not context_text:
                return rag_service_pb2.ChatResponse(
                    answer="No relevant information found in documents.",
                    sources=[],
                )

            system_prompt = (
                "You are a corporate AI assistant. "
                "Answer strictly based on the provided context. "
                "If information is missing, say 'No information found'.\n\n"
                f"Context:\n{context_text}"
            )

            messages = [
                SystemMessage(content=system_prompt),
                HumanMessage(content=query),
            ]

            ai_response = await llm.ainvoke(messages)

            return rag_service_pb2.ChatResponse(
                answer=str(ai_response.content),
                sources=sources_proto,
            )

        except Exception as e:
            logger.exception("RAG error: %s", e)
            return rag_service_pb2.ChatResponse(
                answer="Error processing request.",
                sources=[],
            )


async def serve() -> None:
    """Start gRPC server and RabbitMQ consumer."""
    await init_qdrant_collection()

    server = grpc.aio.server()
    rag_service_pb2_grpc.add_RagServiceServicer_to_server(RagService(), server)
    listen_addr = "[::]:50051"
    server.add_insecure_port(listen_addr)
    logger.info("gRPC server started on %s", listen_addr)

    await server.start()

    try:
        await asyncio.gather(
            server.wait_for_termination(),
            consume_rabbitmq(),
        )
    except asyncio.CancelledError:
        await server.stop(0)


if __name__ == "__main__":
    try:
        asyncio.run(serve())
    except KeyboardInterrupt:
        pass
